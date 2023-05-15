"""
mdfic.docx - Build .docx manuscripts in SFFMS format from html/markdown.

"""

import logging

from math import ceil

from html.parser import HTMLParser
import docx
import zipfile
from docx.shared import Length,Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xml.dom import minidom
import datetime
import re
import os
from .utils import get_in, int_to_roman


############################################################
# DOCX Support


import logging

logger = logging.getLogger(__name__)

METADATA_DEFAULTS = dict(
    title = "Untitled",
    author = 'A. I. Robotsky',
    address = '',
    email = '',
    date = '',
    )


DOCX_HEADER_TEMPLATE = """
<?xml version="1.0" ?><w:hdr mc:Ignorable="w14" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><w:p><w:pPr><w:pStyle w:val="Footnote"/><w:jc w:val="right"/></w:pPr><w:r><w:t>{author} / {title} / </w:t></w:r><w:r><w:fldChar w:fldCharType="begin" w:fldLock="0"/></w:r><w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate" w:fldLock="0"/></w:r><w:r><w:t>11</w:t></w:r><w:r><w:fldChar w:fldCharType="end" w:fldLock="0"/></w:r></w:p></w:hdr>
""".strip()

DOCX_RELS_FILENAME = 'word/_rels/document.xml.rels'
DOCX_DOC_FILENAME = 'word/document.xml'

def set_sffms_styles(d):
    
    for s in d.styles:
        try:
            s.font.name = 'Courier'
            s.font.size = Pt(12)
            s.font.color.rgb = docx.shared.RGBColor(0,0,0)
        except AttributeError:
            logger.warn("Error setting font info for style {s}".format(s=s))

        if s.name == 'Normal':
            s.paragraph_format.line_spacing = 2.0
            s.paragraph_format.first_line_indent = Length(Inches(0.5))
            s.paragraph_format.space_after = Length(Inches(0))
            s.paragraph_format.space_before = Length(Inches(0))
        if s.name == 'Title':
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if s.name == 'List Bullet':
            s.paragraph_format.line_spacing = 1.0
            s.paragraph_format.space_after = Length(Inches(0))
            s.paragraph_format.space_before = Length(Inches(0))
        if re.fullmatch("Heading [1-3]", s.name):
            logger.info("Heading style: {}".format(s.name))
            level = int(s.name[-1])
            if level == 3:
                s.font.all_caps = True
                s.font.underline = False
                s.font.bold = False
            if level == 2:
                s.font.all_caps = True
                s.font.underline = True
                s.font.bold = False
            if level == 1:
                s.font.all_caps = True
                s.font.underline = True
                s.font.bold = True
    
    return input        

def isstrong(tag):
    return tag in ['b','strong']

def isemphasis(tag):
    return tag in ['i','em']

def prettyxml(xml):
    """
    Pretty-print xml
    """
    xml = minidom.parseString(xml_fname) 
    return xml.toprettyxml()

def xml_traverse(xml):
    """
    A generator function that takes parsed xml from the minidom parser
    and yields all the nodes.
    """
    yield xml
    for c in xml.childNodes:
        for n in xml_traverse(c):
            yield n

def xml_get_rel_ids(xml):
    """
    Take the parsed xml from a DOCX relationships file
    (i.e. 'word/_rels/document.xml.rels') and produces all the 
    relationship ids
    """
    result = []
    for node in xml_traverse(xml):
        if node.nodeType == node.ELEMENT_NODE and node.hasAttribute('Id'):
            result.append(node.getAttribute('Id'))
    return result
    
def xml_rel_nums(xml):
    """
    Take the parsed xml from a WORD DOCX relationships file and
    produce all the relationship numbers
    """
    return [int(rel[3:]) for rel in xml_get_rel_ids(xml)]

class HTML2DOCX(HTMLParser):

    def __init__(self,metadata,sffms=False):
        self.metadata = METADATA_DEFAULTS
        self.metadata.update(metadata)
        self.tag_attrs = {}
        self.sffms = sffms
        self.number_scenes = get_in(self.metadata,['mdfic','number_scenes'],False)
        self.scene_number = 1

        super().__init__()

    @property
    def title(self):
        return self.metadata['title']
    @property
    def running_title(self):
        try: 
            return self.metadata['running_title']
        except KeyError:
            return self.title
    @property
    def author(self):
        return self.metadata['author']
    @property
    def address(self):
        return self.metadata['address']
    @property
    def email(self):
        return self.metadata['email']
    @property
    def date(self):
        return self.metadata['date']

    def reset(self):
        super().reset()
        # Add a two column table at the top with the contact info
        # and space for the wordcount
        self.doc = docx.Document()
        self.doc.core_properties.author = self.author
        self.doc.core_properties.title = self.title
        self.doc.core_properties.created = self.doc.core_properties.modified = datetime.datetime.now()
        self.doc.core_properties.category = 'story'

        logger.debug("NUM SECTIONS = {}".format(len(self.doc.sections)))

        self.doc.sections[0].left_margin = Length(Inches(1))
        self.doc.sections[0].right_margin = Length(Inches(1))

        logger.debug("section[0] left margin = {}".format(self.doc.sections[0].left_margin))
        logger.debug("section[0] right margin = {}".format(self.doc.sections[0].right_margin))

        if self.sffms:
            set_sffms_styles(self.doc)
        self.top_table = self.doc.add_table(rows=1,cols=2)
        self.addr_cell = self.top_table.cell(0,0)
        self.wordcount_cell = self.top_table.cell(0,1)
        self.addr_cell.text = '\n'.join(x for x in [self.address, self.email, self.date] if x)
        self.addr_cell.paragraphs[0].style = "No Spacing"

        # Add a spacer
        self.doc.add_paragraph('\n'*8)

        # add the title
        self.doc.add_paragraph(self.title.upper(),  style='Normal').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # add the author
        self.doc.add_paragraph("By " + self.author, style='Normal').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.current_paragraph = self.doc.add_paragraph('', style='Normal')
        self.stack = []
        self.emphasis = 0
        self.strong = 0
        self.blockquote = 0
        self.wordcount = 0
        self.is_closed = False
        
    def save(self,filename):
        # Compute the approximate wordcount and add
        # to the title page header
        txt = "about {} Words".format(int(ceil(self.wordcount/50) * 50))
        self.wordcount_cell.text = txt
        self.wordcount_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add the end marker to the end
        self.current_paragraph = self.doc.add_paragraph(text="# # # # #")
        self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        logger.debug("NUM SECTIONS = {}".format(len(self.doc.sections)))
        logger.debug("section[0] left margin = {}".format(self.doc.sections[0].left_margin))
        logger.debug("section[0] right margin = {}".format(self.doc.sections[0].right_margin))

        self.close()

        if self.sffms:
            # If using SFFMS style, add a header to each page
            tmpfilename = filename + ".tmp"
            logger.info("saving temporary file: {}".format(tmpfilename))
            self.doc.save(tmpfilename)

            try:
                # Add header
                logger.info("adding header")
                with zipfile.ZipFile(tmpfilename,mode='r') as oldzip:
                    with zipfile.ZipFile(filename,mode='w') as newzip:
                        for i,name in enumerate(oldzip.namelist()):
                            if name == '[Content_Types].xml':
                                oldxml = oldzip.read(name)
                                newxml = self.add_header_content_override(oldxml)
                                newzip.writestr(name,newxml,compress_type=8)
                            elif name == DOCX_RELS_FILENAME:
                                oldxml = oldzip.read(name)
                                newxml = self.add_header_rel(oldxml,"rId1000")
                                newzip.writestr(name, newxml, compress_type=8)
                                newzip.writestr("word/header1.xml", self.header_xml(), compress_type=8)
                            elif name == DOCX_DOC_FILENAME:
                                oldxml = oldzip.read(name).decode('utf8')
                                newxml = self.add_header_to_doc(oldxml, "rId1000")
                                newzip.writestr(name, newxml, compress_type=8)
                                #newzip.writestr(name,oldxml, compress_type=8)
                            else:
                                newzip.writestr(name, oldzip.read(name), compress_type=8)
            finally:
                os.remove(tmpfilename)
        else:
            self.doc.save(filename)


    def add_header_content_override(self,contentxmlstr):
        doc = minidom.parseString(contentxmlstr)
        #   <Override ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml" PartName="/word/header1.xml"/>
        elem = doc.createElement("Override")
        elem.setAttribute("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml")
        elem.setAttribute("PartName", "/word/header1.xml")
        doc.childNodes[0].appendChild(elem)
        return doc.toxml()

    def add_header_rel(self,relsxmlstr,rel_id):
        """
        Add the new header relation to the rels xmls
        and return the new xml string
        """
        relsdoc = minidom.parseString(relsxmlstr)
        new_rel = relsdoc.createElement("Relationship")
        new_rel.setAttribute("Id", rel_id)
        new_rel.setAttribute("Target", "header1.xml")
        new_rel.setAttribute("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header")
        relsdoc.childNodes[0].appendChild(new_rel)
        return relsdoc.toxml()

    def add_header_to_doc(self,docxmlstr,rel_id):
        """
        Add the new header reference to the 
        document content.
        """
        docxml = minidom.parseString(docxmlstr)
        # <w:headerReference r:id="rId4" w:type="default"/>
        headerref = docxml.createElement("w:headerReference")
        headerref.setAttribute("r:id", rel_id)
        headerref.setAttribute("w:type", "default")
        for elem  in xml_traverse(docxml):
            if elem.nodeType == minidom.Element.ELEMENT_NODE and elem.tagName == 'w:sectPr':
                sect = elem
        sect.appendChild(headerref)
        sect.appendChild(docxml.createElement("w:titlePg"))
        return docxml.toxml()

    def header_xml(self):

        # see https://social.msdn.microsoft.com/Forums/office/en-US/d52e8532-fc0f-42ce-a40c-55811511d800/how-to-add-header-and-footer-to-docx-file-using-ooxml-format?forum=oxmlsdk
        
        # 1. Add the header as 'word/header1.xml'
        header = DOCX_HEADER_TEMPLATE.format(author=self.author, title=self.running_title.upper())
        headerstr = minidom.parseString(header).toxml()
        return headerstr


    def handle_starttag(self,tag,attrs):

        self.stack.append((tag,dict(attrs)))

        logger.debug("starttag: {tag}. stack = {self.stack}".format(**locals()))

        if isstrong(tag):
            self.strong += 1
        elif isemphasis(tag):
            self.emphasis += 1
        elif tag == 'blockquote':
            self.blockquote += 1
        elif tag == 'p':
            if self.number_scenes and self.scene_number == 1:
                self.insert_scene_break()
            self.current_paragraph = self.doc.add_paragraph(text=None,style="Normal")
            if self.blockquote > 0:
                self.current_paragraph.paragraph_format.left_indent = Length(Inches(1))
                self.current_paragraph.paragraph_format.first_line_indent = Length(Inches(0))
        elif tag == 'li':
            self.current_paragraph = self.doc.add_paragraph(style="List Bullet")
        elif tag == 'hr':
            self.insert_scene_break()
        elif tag == 'a':
            logger.info("Got anchor tag. attrs = {}".format(attrs))
        elif re.fullmatch("[Hh][1-3]",tag):
            self.current_paragraph = self.doc.add_paragraph(style="Heading "+tag[1])

    def insert_scene_break(self):
        if self.number_scenes:
            if self.number_scenes == 'roman':
                num = int_to_roman(self.scene_number)
            else:
                num = str(self.scene_number)
            self.current_paragraph = self.doc.add_paragraph(text=num)
            self.scene_number += 1
        else:
            self.current_paragraph = self.doc.add_paragraph(text="#")
        self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def handle_endtag(self,tag):

        logger.debug("endtag: {tag}. stack = {self.stack}".format(**locals()))

        starttag,attrs = self.stack.pop()
        if tag != starttag:
            raise ValueError("Got </{}>, expected </{}>".format(tag,starttag))

        if isstrong(tag):
            self.strong -= 1
        elif isemphasis(tag):
            self.emphasis -= 1
        elif tag == 'blockquote':
            self.blockquote -= 1
        elif tag == 'a':
            link_text = " ({href}) ".format(**attrs)
            self.current_paragraph.add_run(link_text)


    def handle_data(self,data):

        logger.debug( "----" )
        logger.debug(u"Adding data: '{data}'".format(**locals()))
        logger.debug("stack = {self.stack}".format(**locals()))
        logger.debug("----")

        # skip data that's outside any tags.
        if self.stack:
            run = self.current_paragraph.add_run(data)
            run.bold = self.current_paragraph.style.font.bold or self.strong > 0
            run.italic = self.current_paragraph.style.font.italic or self.emphasis > 0
            self.wordcount += len(data.split())




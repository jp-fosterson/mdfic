"""End-to-end CLI tests for pandoc-required commands.

Skipped automatically by the pandoc-marker hook in conftest.py when pandoc
isn't on PATH.
"""
import zipfile

import pytest
from docx import Document

from mdfic.cli import cli


pytestmark = pytest.mark.pandoc


# latex --------------------------------------------------------

def test_latex_sffms_default_documentclass(cli_runner, single_story):
    result = cli_runner.invoke(cli, ["latex", str(single_story)])
    assert result.exit_code == 0, result.output
    assert "\\documentclass{sffms}" in result.output
    assert "\\title{A Lipsum Day}" in result.output
    assert "\\author{Jane Q. Author}" in result.output
    assert "\\wordcount{" in result.output
    assert "\\begin{document}" in result.output
    assert "\\end{document}" in result.output


def test_latex_sffms_replaces_section_with_textbf(cli_runner, single_story):
    result = cli_runner.invoke(cli, ["latex", str(single_story)])
    assert result.exit_code == 0, result.output
    assert "\\section{" not in result.output
    assert "\\textbf{Lipsum}" in result.output


def test_latex_sffms_scene_numbers_roman(cli_runner, single_story):
    result = cli_runner.invoke(cli, ["latex", str(single_story)])
    assert result.exit_code == 0, result.output
    assert "\\textbf{I}" in result.output
    assert "\\textbf{II}" in result.output


def test_latex_includes_extra_headers(cli_runner, single_story):
    result = cli_runner.invoke(cli, ["latex", str(single_story)])
    assert result.exit_code == 0, result.output
    assert "\\usepackage{microtype}" in result.output


def test_latex_article_documentclass(cli_runner, single_story):
    result = cli_runner.invoke(cli, ["latex", "--documentclass=article", str(single_story)])
    assert result.exit_code == 0, result.output
    assert "\\documentclass{article}" in result.output
    assert "\\maketitle" in result.output
    assert "\\section{Lipsum}" in result.output


def test_latex_multi_file(cli_runner, multi_metadata, multi_parts):
    args = ["latex", str(multi_metadata)] + [str(p) for p in multi_parts]
    result = cli_runner.invoke(cli, args)
    assert result.exit_code == 0, result.output
    assert "\\title{Two-Part Lipsum}" in result.output
    assert "\\textbf{Part One}" in result.output
    assert "\\textbf{Part Two}" in result.output


# docx ---------------------------------------------------------

def test_docx_plain_single(cli_runner, single_story, tmp_path):
    out = tmp_path / "story.docx"
    result = cli_runner.invoke(
        cli,
        ["docx", "--no-sffms", "--no-date", "-o", str(out), str(single_story)],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    assert out.stat().st_size > 0
    doc = Document(str(out))
    paragraphs = [p.text for p in doc.paragraphs]
    assert "A LIPSUM DAY" in paragraphs
    assert any("Jane Q. Author" in p for p in paragraphs)
    assert "# # # # #" in paragraphs


def test_docx_sffms_injects_header(cli_runner, single_story, tmp_path):
    out = tmp_path / "story.docx"
    result = cli_runner.invoke(
        cli,
        ["docx", "--sffms", "--no-date", "-o", str(out), str(single_story)],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    with zipfile.ZipFile(out) as z:
        assert "word/header1.xml" in z.namelist()


def test_docx_multi(cli_runner, multi_metadata, multi_parts, tmp_path):
    out = tmp_path / "story.docx"
    args = [
        "docx", "--no-sffms", "--no-date", "-o", str(out), str(multi_metadata),
    ] + [str(p) for p in multi_parts]
    result = cli_runner.invoke(cli, args)
    assert result.exit_code == 0, result.output
    doc = Document(str(out))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "TWO-PART LIPSUM" in body
    assert "Part One" in body
    assert "Part Two" in body


# html ---------------------------------------------------------

def test_html_single(cli_runner, single_story, tmp_path):
    out = tmp_path / "story.html"
    result = cli_runner.invoke(cli, ["html", "-o", str(out), str(single_story)])
    assert result.exit_code == 0, result.output
    content = out.read_text()
    assert "<html" in content
    assert "Lorem ipsum" in content
    assert "<bold>END</bold>" in content


def test_html_scene_numbers_roman(cli_runner, single_story, tmp_path):
    out = tmp_path / "story.html"
    result = cli_runner.invoke(cli, ["html", "-o", str(out), str(single_story)])
    assert result.exit_code == 0, result.output
    content = out.read_text()
    assert "<bold>I</bold>" in content
    assert "<bold>II</bold>" in content


def test_html_multi(cli_runner, multi_metadata, multi_parts, tmp_path):
    out = tmp_path / "story.html"
    args = ["html", "-o", str(out), str(multi_metadata)] + [str(p) for p in multi_parts]
    result = cli_runner.invoke(cli, args)
    assert result.exit_code == 0, result.output
    content = out.read_text()
    assert "Part One" in content
    assert "Part Two" in content
